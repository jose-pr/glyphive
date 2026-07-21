from __future__ import annotations

import pytest


def _make_frame(codec, kind, idx, payload):
    """Build a single CRC-valid L/P frame line for the given index/payload."""
    from glyphive.codec.engine import _check_chars, encode_index, nibble_encode

    token = encode_index(idx)
    body = nibble_encode(payload)
    return f"{kind}{token} {body} #{_check_chars(kind, token, body)}"


def test_merge_ocr_lines_unions_valid_frames_across_passes():
    """Different blur passes contribute different CRC-valid lines; merge unions them.

    Real-world scan finding (2026-07-17): different Gaussian blur radii recover
    different lines, and the per-line CRC makes combining passes safe. The merge
    keeps the first pass verbatim (line order drives page attribution) and only
    appends CRC-valid frames later passes read that the first pass missed.
    """
    from glyphive import codec
    from glyphive.cli._common import _merge_ocr_lines

    c = codec.get("base16g-crc16-rs")
    good0 = _make_frame(c, "L", 0, b"\x01\x02")
    good1 = _make_frame(c, "L", 1, b"\x03\x04")
    good2 = _make_frame(c, "L", 2, b"\x05\x06")
    garbage = "L000ZZ notavalidframe"  # unparseable, must be ignored

    # Pass A reads frames 0 and 1; pass B reads 0 (dup) and 2 (new).
    merged = _merge_ocr_lines([[good0, good1, garbage], [good0, good2]])

    # First pass preserved verbatim as the ordered spine.
    assert merged[: 3] == [good0, good1, garbage]
    # The new CRC-valid frame from pass B is appended; the duplicate is not.
    assert good2 in merged
    assert merged.count(good0) == 1
    assert merged.count(good2) == 1


def test_merge_ocr_lines_single_pass_is_verbatim():
    from glyphive.cli._common import _merge_ocr_lines

    lines = ["#!glyphive header", "LMYCVH payload #ABCD", "junk"]
    assert _merge_ocr_lines([lines]) == lines


def _write_real_png(path):
    """Write a tiny but genuine PNG -- blur radii > 0 open the file via PIL,
    so a fake magic-bytes-only stub (fine for magic-sniffing tests) fails
    there with ``UnidentifiedImageError``."""
    from PIL import Image

    Image.new("RGB", (4, 4), color="white").save(str(path), format="PNG")


def test_image_ocr_passes_spine_avoids_reocr_of_its_own_radius(tmp_path, monkeypatch):
    """A ``spine`` is merged in without OCR'ing it again -- only ``blur``'s
    radii are OCR'd. Regression for the auto-descan retry: the sharp (0.0)
    pass must not be repeated when it is threaded back in as the spine.
    """
    from glyphive import codec
    from glyphive.cli._common import _image_ocr_passes
    from glyphive.restore import ocr

    image = tmp_path / "photo.png"
    _write_real_png(image)

    c = codec.get("base16g-crc16-rs")
    spine_frame = _make_frame(c, "L", 0, b"\x01\x02")
    extra_frame = _make_frame(c, "L", 1, b"\x03\x04")

    ocr_calls = []

    def counting_ocr_pages(paths, *, engine=None):
        ocr_calls.append(len(paths))
        return [[extra_frame]]

    monkeypatch.setattr(ocr, "ocr_pages", counting_ocr_pages)

    result = _image_ocr_passes(image, blur=[0.6, 0.8], spine=[spine_frame])

    # OCR ran exactly once per NEW radius (0.6, 0.8) -- never for the spine's
    # own (0.0) pass, which was supplied pre-computed instead of re-OCR'd.
    assert len(ocr_calls) == 2
    texts = [line.text for line in result]
    # Spine kept verbatim first, then the new CRC-valid frame appended.
    assert texts == [spine_frame, extra_frame]


def test_image_ocr_passes_spine_matches_prior_full_sweep_behavior(tmp_path, monkeypatch):
    """The spine-based retry produces the SAME merged lines, in the SAME
    order, as the old approach of re-OCRing radius 0.0 alongside 0.6/0.8 --
    only the redundant extra OCR call is removed, not the recovered data.
    """
    from glyphive import codec
    from glyphive.cli._common import _image_ocr_passes
    from glyphive.restore import ocr

    image = tmp_path / "photo.png"
    _write_real_png(image)

    c = codec.get("base16g-crc16-rs")
    sharp_frame = _make_frame(c, "L", 0, b"\x01\x02")
    blur_06_frame = _make_frame(c, "L", 1, b"\x03\x04")
    blur_08_frame = _make_frame(c, "L", 2, b"\x05\x06")

    per_radius = {0.0: [sharp_frame], 0.6: [blur_06_frame], 0.8: [blur_08_frame]}

    # --- Old behavior: full sweep over [0.0, 0.6, 0.8], OCRing 0.0 again. ---
    def ocr_by_call_order_full(paths, *, engine=None, _order=iter([0.0, 0.6, 0.8])):
        return [[line for line in per_radius[next(_order)]]]

    monkeypatch.setattr(ocr, "ocr_pages", ocr_by_call_order_full)
    old_result = [
        line.text for line in _image_ocr_passes(image, blur=[0.0, 0.6, 0.8])
    ]

    # --- New behavior: sharp pass precomputed, retry only OCRs 0.6/0.8. ---
    def ocr_by_call_order_extra(paths, *, engine=None, _order=iter([0.6, 0.8])):
        return [[line for line in per_radius[next(_order)]]]

    monkeypatch.setattr(ocr, "ocr_pages", ocr_by_call_order_extra)
    new_result = [
        line.text
        for line in _image_ocr_passes(image, blur=[0.6, 0.8], spine=[sharp_frame])
    ]

    assert new_result == old_result


def test_auto_input_renders_pdf_then_ocr(tmp_path, monkeypatch):
    from glyphive.cli import _common
    from glyphive.restore import ocr

    source = tmp_path / "scan.pdf"
    source.write_bytes(b"pdf")
    seen = {}

    def fake_render(path, destination, *, blur=0.0):
        seen["render"] = path.name
        return [destination / "scan-0001.png", destination / "scan-0002.png"]

    def fake_ocr(paths, *, engine=None):
        seen.setdefault("ocr", []).append(([path.name for path in paths], engine))
        return [[f"line-{paths[0].name}"]]

    monkeypatch.setattr(
        "glyphive.restore.document_images.render_document_images", fake_render
    )
    monkeypatch.setattr(ocr, "ocr_pages", fake_ocr)

    assert _common.load_input_lines(source, engine="mock") == [
        "line-scan-0001.png"
    ]
    assert seen == {
        "render": "scan.pdf",
        "ocr": [(["scan-0001.png", "scan-0002.png"], "mock")],
    }


def test_auto_directory_accepts_mixed_supported_inputs(tmp_path, monkeypatch):
    from glyphive.cli import _common
    from glyphive.restore import ocr

    inputs = tmp_path / "inputs"
    inputs.mkdir()
    (inputs / "a.txt").write_text("transcript\n", encoding="utf-8")
    (inputs / "b.png").write_bytes(b"png")
    monkeypatch.setattr(
        ocr,
        "ocr_pages",
        lambda paths, *, engine=None: [[f"ocr-{paths[0].name}"]],
    )
    assert _common.load_input_lines(inputs) == ["transcript", "ocr-b.png"]


def test_image_magic_wins_over_missing_or_wrong_extension(tmp_path, monkeypatch):
    from glyphive.cli import _common
    from glyphive.restore import ocr

    image = tmp_path / "renamed.txt"
    image.write_bytes(b"\x89PNG\r\n\x1a\nnot-a-real-image")
    monkeypatch.setattr(
        ocr, "ocr_pages", lambda paths, *, engine=None: [["ocr-page"]]
    )
    assert _common.load_input_lines(image) == ["ocr-page"]


def test_pdf_magic_wins_over_extension(tmp_path, monkeypatch):
    from glyphive.cli import _common
    from glyphive.restore import ocr

    pdf = tmp_path / "scan.bin"
    pdf.write_bytes(b"%PDF-1.7\n")
    monkeypatch.setattr(
        "glyphive.restore.document_images.render_document_images",
        lambda path, destination, *, blur=0.0: [destination / "page.png"],
    )
    monkeypatch.setattr(ocr, "ocr_pages", lambda paths, *, engine=None: [["pdf-page"]])
    assert _common.load_input_lines(pdf) == ["pdf-page"]


def test_unknown_binary_input_fails_clearly(tmp_path):
    from glyphive.cli._common import load_input_lines

    source = tmp_path / "unknown.bin"
    source.write_bytes(b"\x00\xff\x00\xff")
    with pytest.raises(ValueError, match="cannot detect supported input type"):
        load_input_lines(source)


def test_list_uses_auto_input_and_forwards_ocr_engine(tmp_path, monkeypatch):
    from glyphive import cli
    from glyphive.cli import list as list_command
    from glyphive.restore import decode

    seen = {}
    monkeypatch.setattr(
        list_command,
        "load_input_lines",
        lambda source, engine=None, blur=0.0, spine=None: seen.update(
            source=source, engine=engine
        ) or [],
    )
    monkeypatch.setattr(
        decode,
        "decode_document_to_spool",
        lambda lines, sink, **options: {
                "v": 1,
                "codec": "base16g-crc16-rs",
                "comp": "none",
                "files": 0,
                "bytes": 0,
                "pages": 1,
            },
    )
    monkeypatch.setattr(list_command._archive, "iter_record_events", lambda raw, **options: [])

    assert cli.run(
        ["list", "-f", str(tmp_path / "scan.pdf"), "--ocr-engine", "mock"]
    ) == 0
    assert seen["source"].name == "scan.pdf"
    assert seen["engine"] == "mock"


def test_docx_transcript_is_read_directly_and_diagnostic_pages_render(tmp_path):
    import docx
    from docx.enum.text import WD_BREAK

    from glyphive.cli._common import load_input_lines
    from glyphive.restore.document_images import render_document_images

    source = tmp_path / "scan.docx"
    document = docx.Document()
    document.add_paragraph("first")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("second")
    document.save(source)

    assert load_input_lines(source) == ["first", "second"]
    outputs = render_document_images(source, tmp_path / "pages", dpi=72)
    assert [path.name for path in outputs] == ["scan-0001.png", "scan-0002.png"]
    assert all(path.read_bytes().startswith(b"\x89PNG") for path in outputs)


@pytest.mark.parametrize("dpi,blur", [(0, 0), (300, -1)])
def test_document_render_options_are_validated(tmp_path, dpi, blur):
    from glyphive.restore.document_images import render_document_images

    with pytest.raises(ValueError):
        render_document_images(
            tmp_path / "scan.pdf", tmp_path / "pages", dpi=dpi, blur=blur
        )
