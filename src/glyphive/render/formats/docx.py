"""Word renderer with a lazy python-docx import."""

from __future__ import annotations

import os as _os
import typing as _ty

from pathlib_next import Path

from glyphive.layout import Page
from glyphive.render._base import (
    DEFAULT_DOCX_FONT,
    DEFAULT_PAGE_MARGIN_PT,
    HORIZONTAL_ALIGNMENTS,
    RenderFormat,
)


class DocxRenderFormat(RenderFormat):
    name = "docx"

    @classmethod
    def is_available(cls) -> bool:
        try:
            import importlib.util

            return importlib.util.find_spec("docx") is not None
        except Exception:
            return False

    def render(
        self,
        pages: _ty.Iterable[Page],
        out: _ty.Union[str, "_os.PathLike[str]"],
        *,
        font: _ty.Optional[str] = None,
        font_size: float = 11.0,
        page_margin_pt: float = DEFAULT_PAGE_MARGIN_PT,
        horizontal_alignment: str = "left",
        character_spacing_pt: float = 0.0,
    ) -> None:
        try:
            import docx
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except ImportError as exc:
            raise RuntimeError(
                "Word output needs the 'python-docx' backend; install glyphive[docx]"
            ) from exc
        if font_size <= 0:
            raise ValueError("font_size must be > 0")
        if page_margin_pt < 0 or page_margin_pt * 2 >= 612.0:
            raise ValueError("page_margin_pt must leave positive printable width")
        if horizontal_alignment not in HORIZONTAL_ALIGNMENTS:
            raise ValueError("horizontal_alignment must be left, center, or justify")
        if character_spacing_pt < 0:
            raise ValueError("character_spacing_pt must be >= 0")
        family = font or DEFAULT_DOCX_FONT
        size = Pt(font_size)
        document = docx.Document()
        for section in document.sections:
            section.top_margin = Pt(page_margin_pt)
            section.bottom_margin = Pt(page_margin_pt)
            section.left_margin = Pt(page_margin_pt)
            section.right_margin = Pt(page_margin_pt)
        normal = document.styles["Normal"]
        normal.font.name = family
        normal.font.size = size
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), family)
        for page_index, page in enumerate(pages):
            if page_index > 0:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            for line in page.text_lines:
                para = document.add_paragraph()
                fmt = para.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1.0
                fmt.alignment = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "justify": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                }[horizontal_alignment]
                run = para.add_run(line)
                run.font.name = family
                run.font.size = size
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                for slot in ("ascii", "hAnsi", "cs", "eastAsia"):
                    rfonts.set(qn(f"w:{slot}"), family)
                if character_spacing_pt:
                    spacing = OxmlElement("w:spacing")
                    spacing.set(qn("w:val"), str(round(character_spacing_pt * 20)))
                    rpr.append(spacing)
        document.save(str(Path(_os.fspath(out))))
