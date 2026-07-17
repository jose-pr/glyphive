"""Render document pages to images for OCR and troubleshooting."""

from __future__ import annotations

import typing as _ty
import zipfile

from pathlib_next import Path

__all__ = ["read_docx_lines", "render_document_images"]


def render_document_images(
    source: _ty.Union[str, "Path"],
    destination: _ty.Union[str, "Path"],
    *,
    dpi: int = 300,
    blur: float = 0.0,
) -> _ty.List["Path"]:
    """Render a PDF or DOCX to ordered PNG files.

    PDF pages preserve their physical layout. DOCX pages are a diagnostic
    re-render of paragraph text because python-docx does not implement Word's
    layout engine; direct Glyphive restore reads their transcript text instead.
    """
    source_path = Path(source)
    destination_path = Path(destination)
    if dpi <= 0:
        raise ValueError("dpi must be greater than zero")
    if blur < 0:
        raise ValueError("blur must be zero or greater")
    kind = _document_kind(source_path)
    if kind == "pdf":
        return _render_pdf(source_path, destination_path, dpi=dpi, blur=blur)
    if kind == "docx":
        return _render_docx(source_path, destination_path, dpi=dpi, blur=blur)
    raise ValueError(
        f"unsupported document input {source_path}; expected .pdf or .docx"
    )


def _document_kind(source: "Path") -> str:
    with source.open("rb") as stream:
        prefix = stream.read(8)
    if prefix.startswith(b"%PDF-"):
        return "pdf"
    if prefix.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(str(source)) as archive:
                if "word/document.xml" in archive.namelist():
                    return "docx"
        except (OSError, zipfile.BadZipFile):
            pass
    suffix = source.suffix.lower()
    return suffix[1:] if suffix in {".docx", ".pdf"} else "unknown"


def _read_docx_pages(source: "Path") -> _ty.List[_ty.List[str]]:
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise RuntimeError(
            "DOCX input requires python-docx; install glyphive[docx]"
        ) from exc

    document = docx.Document(str(source))
    pages: _ty.List[_ty.List[str]] = [[]]
    for paragraph in document.paragraphs:
        has_page_break = any(
            node.get(qn("w:type")) == "page"
            for node in paragraph._p.iter(qn("w:br"))
        )
        if has_page_break and pages[-1]:
            pages.append([])
        if paragraph.text:
            pages[-1].append(paragraph.text)
    return [page for page in pages if page]


def read_docx_lines(source: _ty.Union[str, "Path"]) -> _ty.List[str]:
    """Read paragraph lines from a Glyphive-generated DOCX without OCR."""
    pages = _read_docx_pages(Path(source))
    lines = [line for page in pages for line in page]
    if not lines:
        raise ValueError("DOCX contains no paragraph transcript to restore")
    return lines


def _render_docx(
    source: "Path", destination: "Path", *, dpi: int, blur: float
) -> _ty.List["Path"]:
    try:
        from importlib import resources

        from PIL import Image, ImageDraw, ImageFilter, ImageFont
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise RuntimeError(
            "DOCX diagnostic rendering requires python-docx and Pillow; install "
            "glyphive[docx,document-input]"
        ) from exc

    pages = _read_docx_pages(source)
    if not pages:
        raise ValueError("DOCX contains no paragraph transcript to render")
    destination.mkdir(parents=True, exist_ok=True)
    font_resource = resources.files("glyphive.assets.fonts.ocr_b").joinpath(
        "OCR-B.ttf"
    )
    font = ImageFont.truetype(str(font_resource), max(1, round(8 * dpi / 72)))
    width, height = round(8.5 * dpi), round(11 * dpi)
    margin = round(0.5 * dpi)
    leading = max(1, round(8 * 1.2 * dpi / 72))
    outputs: _ty.List["Path"] = []
    for index, page in enumerate(pages, 1):
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        y = margin
        for line in page:
            draw.text((margin, y), line, font=font, fill="black")
            y += leading
        if blur:
            image = image.filter(ImageFilter.GaussianBlur(radius=blur))
        output = destination / f"{source.stem}-{index:04d}.png"
        image.save(str(output), format="PNG", dpi=(dpi, dpi))
        outputs.append(output)
    return outputs


def _render_pdf(
    source: "Path", destination: "Path", *, dpi: int, blur: float
) -> _ty.List["Path"]:
    try:
        import pypdfium2
        from PIL import ImageFilter
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise RuntimeError(
            "PDF page rendering requires pypdfium2 and Pillow; install "
            "glyphive[document-input]"
        ) from exc
    destination.mkdir(parents=True, exist_ok=True)
    document = pypdfium2.PdfDocument(str(source))
    outputs: _ty.List["Path"] = []
    try:
        for index in range(len(document)):
            image = document[index].render(scale=dpi / 72).to_pil().convert("RGB")
            if blur:
                image = image.filter(ImageFilter.GaussianBlur(radius=blur))
            output = destination / f"{source.stem}-{index + 1:04d}.png"
            image.save(str(output), format="PNG", dpi=(dpi, dpi))
            outputs.append(output)
    finally:
        document.close()
    return outputs
